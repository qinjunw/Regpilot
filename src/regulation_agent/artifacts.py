from __future__ import annotations

import hashlib
import json
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

from .source_documents import DocumentSourceStore


SUPPORTED_REPORT_FORMATS = {"md", "docx"}
MAX_REPORT_MARKDOWN_CHARS = 240_000


class InterpretationArtifactStore:
    def __init__(self, state_dir: str | Path, source_store: DocumentSourceStore) -> None:
        self.state_dir = Path(state_dir).expanduser()
        self.source_store = source_store
        self.default_output_dir = self.state_dir / "artifacts" / "interpretation_reports"
        self.default_output_dir.mkdir(parents=True, exist_ok=True)

    def generate_interpretation_report(
        self,
        *,
        collection_id: str,
        title: str,
        markdown: str,
        source_evidence_ids: list[str],
        formats: list[str] | None = None,
        output_dir: str | Path | None = None,
        filename: str = "",
        overwrite: bool = False,
    ) -> dict[str, Any]:
        clean_title = _single_line(title) or "法规解读报告"
        clean_markdown = _clean_markdown(markdown)
        if not clean_markdown:
            raise ValueError("Report markdown is required.")
        if len(clean_markdown) > MAX_REPORT_MARKDOWN_CHARS:
            raise ValueError("Report markdown exceeds the current size limit.")
        clean_formats = _clean_formats(formats)
        collection = self.source_store.read_collection(collection_id)
        evidence_records = {
            str(item.get("evidence_id") or ""): item
            for item in collection.get("evidence", [])
            if str(item.get("evidence_id") or "").strip()
        }
        known_evidence = set(evidence_records)
        clean_evidence_ids = [str(item or "").strip() for item in source_evidence_ids if str(item or "").strip()]
        if not clean_evidence_ids:
            raise ValueError("source_evidence_ids must include at least one Source Evidence id used by the report.")
        unknown = [item for item in clean_evidence_ids if item not in known_evidence]
        if unknown:
            raise ValueError(f"source_evidence_ids are not in the source collection: {', '.join(unknown[:3])}")
        locator_reference_count = _locator_reference_count(clean_markdown, clean_evidence_ids, evidence_records)
        if locator_reference_count <= 0:
            raise ValueError("Report markdown must cite at least one provided Source Evidence id or locator.")
        target_dir = Path(output_dir).expanduser() if output_dir else self.default_output_dir
        target_dir.mkdir(parents=True, exist_ok=True)
        base_name = _safe_filename(filename or clean_title)
        artifact_id = "artifact_" + hashlib.sha256(
            f"{collection.get('collection_id')}|{clean_title}|{clean_markdown}".encode("utf-8")
        ).hexdigest()[:16]
        artifact_paths = []
        candidate_paths = [_artifact_path(target_dir, base_name, fmt) for fmt in clean_formats]
        manifest_path = target_dir / f"{base_name}.regpilot.json"
        for path in [*candidate_paths, manifest_path]:
            if path.exists() and not overwrite:
                raise ValueError(f"Artifact already exists: {path}")
        format_quality: dict[str, Any] = {}
        for fmt, path in zip(clean_formats, candidate_paths):
            if fmt == "md":
                path.write_text(clean_markdown, encoding="utf-8")
            elif fmt == "docx":
                format_quality["docx"] = _write_docx_from_markdown(path, clean_title, clean_markdown)
            artifact_paths.append({"format": fmt, "path": str(path)})
        quality = {
            "source_evidence_count": len(clean_evidence_ids),
            "locator_reference_count": locator_reference_count,
            "docx": format_quality.get("docx", {}),
        }
        manifest = {
            "artifact_id": artifact_id,
            "title": clean_title,
            "collection_id": str(collection.get("collection_id") or ""),
            "collection_name": str(collection.get("collection_name") or ""),
            "source_evidence_ids": clean_evidence_ids,
            "formats": clean_formats,
            "artifacts": artifact_paths,
            "quality": quality,
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        return {
            "ok": True,
            "artifact_id": artifact_id,
            "title": clean_title,
            "collection_id": str(collection.get("collection_id") or ""),
            "source_evidence_ids": clean_evidence_ids,
            "artifacts": artifact_paths,
            "manifest_path": str(manifest_path),
            "quality": quality,
            "message": f"已生成 {len(artifact_paths)} 个法规解读文件。",
        }


def _clean_formats(formats: list[str] | None) -> list[str]:
    requested = formats or ["md", "docx"]
    result = []
    for item in requested:
        fmt = str(item or "").strip().lower().lstrip(".")
        if fmt not in SUPPORTED_REPORT_FORMATS:
            raise ValueError(f"Unsupported report format: {item}")
        if fmt not in result:
            result.append(fmt)
    return result or ["md", "docx"]


def _clean_markdown(markdown: str) -> str:
    text = str(markdown or "").replace("\r\n", "\n").replace("\r", "\n").strip()
    return text + "\n" if text else ""


def _artifact_path(directory: Path, base_name: str, fmt: str) -> Path:
    return directory / f"{base_name}.{fmt}"


def _safe_filename(value: str) -> str:
    text = _single_line(value).lower()
    text = re.sub(r"[^\w\u4e00-\u9fff.-]+", "-", text, flags=re.UNICODE)
    text = re.sub(r"-{2,}", "-", text).strip(" .-_")
    return text[:80] or "interpretation-report"


def _single_line(value: str) -> str:
    return " ".join(str(value or "").split())[:160]


def _locator_reference_count(markdown: str, evidence_ids: list[str], evidence_records: dict[str, dict[str, Any]]) -> int:
    haystack = str(markdown or "").casefold()
    count = 0
    seen: set[str] = set()
    for evidence_id in evidence_ids:
        if evidence_id in seen:
            continue
        seen.add(evidence_id)
        evidence = evidence_records.get(evidence_id)
        locator = evidence.get("locator") if isinstance(evidence, dict) and isinstance(evidence.get("locator"), dict) else {}
        tokens = [evidence_id, *(_locator_tokens(locator))]
        if any(token and token.casefold() in haystack for token in tokens):
            count += 1
    return count


def _locator_tokens(locator: dict[str, Any]) -> list[str]:
    locator_type = str(locator.get("type") or "")
    if locator_type == "line":
        start = int(locator.get("start_line") or 0)
        end = int(locator.get("end_line") or start)
        return [f"line {start}", f"lines {start}-{end}", f"第{start}行"]
    if locator_type == "docx_paragraph":
        paragraph = str(locator.get("paragraph") or "")
        return [f"DOCX paragraph {paragraph}", f"paragraph {paragraph}", f"第{paragraph}段"]
    if locator_type == "xlsx_row":
        sheet = str(locator.get("sheet") or "Sheet")
        row = str(locator.get("row") or "")
        return [f"XLSX {sheet} R{row}", f"{sheet} R{row}", f"{sheet} 第{row}行"]
    if locator_type == "pdf_page":
        page = str(locator.get("page") or "")
        return [f"PDF p.{page}", f"PDF page {page}", f"p.{page}", f"第{page}页"]
    if locator_type == "pdf_text_object":
        return ["PDF text object"]
    return []


def _write_docx_from_markdown(path: Path, title: str, markdown: str) -> dict[str, Any]:
    try:
        return _write_docx_with_python_docx(path, title, markdown)
    except ImportError:
        _write_docx_fallback_ooxml(path, title, markdown)
        return {"renderer": "minimal_ooxml", "table_count": 0, "footer": False, "paragraph_count": len(_markdown_blocks(markdown))}


def _write_docx_with_python_docx(path: Path, title: str, markdown: str) -> dict[str, Any]:
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    _configure_docx_styles(document, qn, Pt, RGBColor)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_paragraph.add_run(title)
    _set_run_font(title_run, qn, "Microsoft YaHei", Pt(18), bold=True, color=RGBColor(31, 78, 121))
    title_paragraph.paragraph_format.space_after = Pt(10)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(f"RegPilot Source Evidence Artifact | {datetime.now(timezone.utc).date().isoformat()}")
    _set_run_font(footer_run, qn, "Microsoft YaHei", Pt(8), color=RGBColor(96, 96, 96))

    table_count = 0
    paragraph_count = 1
    for block in _markdown_blocks(markdown):
        kind = block["kind"]
        if kind == "table":
            if _add_docx_table(document, block["rows"], qn, OxmlElement, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, Pt, RGBColor):
                table_count += 1
            continue
        text = block["text"]
        if kind == "heading1":
            paragraph = document.add_heading(text, level=1)
        elif kind == "heading2":
            paragraph = document.add_heading(text, level=2)
        elif kind == "heading3":
            paragraph = document.add_heading(text, level=3)
        elif kind == "bullet":
            paragraph = document.add_paragraph(text, style="List Bullet")
        elif kind == "number":
            paragraph = document.add_paragraph(text, style="List Number")
        else:
            paragraph = document.add_paragraph(text)
        paragraph_count += 1
        paragraph.paragraph_format.space_after = Pt(6)

    document.save(path)
    return {"renderer": "python-docx", "table_count": table_count, "footer": True, "paragraph_count": paragraph_count}


def _configure_docx_styles(document: Any, qn: Any, Pt: Any, RGBColor: Any) -> None:
    _set_style_font(document.styles["Normal"], qn, "Microsoft YaHei", Pt(10.5))
    for style_name, size, color in (
        ("Heading 1", 15, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 11.5, RGBColor(31, 78, 121)),
    ):
        style = document.styles[style_name]
        _set_style_font(style, qn, "Microsoft YaHei", Pt(size), bold=True, color=color)


def _set_style_font(style: Any, qn: Any, font_name: str, size: Any, *, bold: bool = False, color: Any | None = None) -> None:
    style.font.name = font_name
    style.font.size = size
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        from docx.oxml import OxmlElement

        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _set_run_font(run: Any, qn: Any, font_name: str, size: Any, *, bold: bool = False, color: Any | None = None) -> None:
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        from docx.oxml import OxmlElement

        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def _add_docx_table(
    document: Any,
    rows: list[list[str]],
    qn: Any,
    OxmlElement: Any,
    WD_CELL_VERTICAL_ALIGNMENT: Any,
    WD_TABLE_ALIGNMENT: Any,
    Pt: Any,
    RGBColor: Any,
) -> bool:
    clean_rows = [[_strip_inline_markdown(cell) for cell in row] for row in rows if any(str(cell).strip() for cell in row)]
    if not clean_rows:
        return False
    column_count = max(len(row) for row in clean_rows)
    table = document.add_table(rows=len(clean_rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    for row_index, row in enumerate(clean_rows):
        table_row = table.rows[row_index]
        if row_index == 0:
            tr_pr = table_row._tr.get_or_add_trPr()
            tr_pr.append(OxmlElement("w:tblHeader"))
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[column_index] if column_index < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            _set_run_font(run, qn, "Microsoft YaHei", Pt(9.5), bold=row_index == 0)
            if row_index == 0:
                run.font.color.rgb = RGBColor(31, 78, 121)
                tc_pr = cell._tc.get_or_add_tcPr()
                shade = OxmlElement("w:shd")
                shade.set(qn("w:fill"), "D9EAF7")
                tc_pr.append(shade)
    document.add_paragraph()
    return True


def _write_docx_fallback_ooxml(path: Path, title: str, markdown: str) -> None:
    document_xml = _document_xml(title, markdown)
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", _content_types_xml())
        archive.writestr("_rels/.rels", _root_relationships_xml())
        archive.writestr("docProps/core.xml", _core_properties_xml(title))
        archive.writestr("word/_rels/document.xml.rels", _document_relationships_xml())
        archive.writestr("word/styles.xml", _styles_xml())
        archive.writestr("word/document.xml", document_xml)


def _document_xml(title: str, markdown: str) -> str:
    paragraphs = [_paragraph_xml(title, style="Title")]
    for block in _markdown_blocks(markdown):
        kind = block["kind"]
        if kind == "table":
            paragraphs.append(_table_xml(block["rows"]))
        elif kind == "heading1":
            text = block["text"]
            paragraphs.append(_paragraph_xml(text, style="Heading1"))
        elif kind == "heading2":
            text = block["text"]
            paragraphs.append(_paragraph_xml(text, style="Heading2"))
        elif kind == "heading3":
            text = block["text"]
            paragraphs.append(_paragraph_xml(text, style="Heading3"))
        elif kind == "bullet":
            text = block["text"]
            paragraphs.append(_paragraph_xml(text, style="ListParagraph", prefix="- "))
        elif kind == "number":
            text = block["text"]
            paragraphs.append(_paragraph_xml(text, style="ListParagraph", prefix=f"{block['number']}. "))
        else:
            text = block["text"]
            paragraphs.append(_paragraph_xml(text, style="Normal"))
    body = "".join(paragraphs)
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{body}<w:sectPr><w:pgSz w:w=\"11906\" w:h=\"16838\"/><w:pgMar w:top=\"1440\" w:right=\"1440\" w:bottom=\"1440\" w:left=\"1440\" w:header=\"720\" w:footer=\"720\" w:gutter=\"0\"/></w:sectPr></w:body>"
        "</w:document>"
    )


def _markdown_blocks(markdown: str) -> list[dict[str, Any]]:
    blocks = []
    pending: list[str] = []

    def flush() -> None:
        if pending:
            blocks.append({"kind": "paragraph", "text": " ".join(pending).strip()})
            pending.clear()

    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        if not line:
            flush()
            index += 1
            continue
        if _looks_like_table_start(lines, index):
            flush()
            table_lines = []
            while index < len(lines) and _is_table_row(lines[index].strip()):
                table_lines.append(lines[index].strip())
                index += 1
            rows = [_parse_table_row(item) for item in table_lines if not _is_table_separator_row(item)]
            if rows:
                blocks.append({"kind": "table", "rows": rows})
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            flush()
            level = len(heading.group(1))
            blocks.append({"kind": f"heading{level}", "text": _strip_inline_markdown(heading.group(2))})
            index += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            flush()
            blocks.append({"kind": "bullet", "text": _strip_inline_markdown(bullet.group(1))})
            index += 1
            continue
        number = re.match(r"^(\d+)[.)]\s+(.+)$", line)
        if number:
            flush()
            blocks.append({"kind": "number", "number": number.group(1), "text": _strip_inline_markdown(number.group(2))})
            index += 1
            continue
        pending.append(_strip_inline_markdown(line))
        index += 1
    flush()
    return blocks


def _looks_like_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    return _is_table_row(lines[index].strip()) and _is_table_separator_row(lines[index + 1].strip())


def _is_table_row(line: str) -> bool:
    return line.count("|") >= 2 and line.startswith("|") and line.endswith("|")


def _is_table_separator_row(line: str) -> bool:
    if not _is_table_row(line):
        return False
    cells = _parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _parse_table_row(line: str) -> list[str]:
    return [_strip_inline_markdown(cell.strip()) for cell in line.strip().strip("|").split("|")]


def _strip_inline_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def _paragraph_xml(text: str, *, style: str, prefix: str = "") -> str:
    style_xml = "" if style == "Normal" else f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>'
    runs = "".join(_run_xml(part) for part in (prefix + text).split("\n"))
    return f"<w:p>{style_xml}{runs}</w:p>"


def _table_xml(rows: list[list[str]]) -> str:
    clean_rows = [[str(cell or "") for cell in row] for row in rows if any(str(cell or "").strip() for cell in row)]
    if not clean_rows:
        return ""
    column_count = max(len(row) for row in clean_rows)
    grid = "".join('<w:gridCol w:w="2400"/>' for _ in range(column_count))
    row_xml = []
    for row in clean_rows:
        cells = []
        for index in range(column_count):
            text = row[index] if index < len(row) else ""
            cells.append(
                '<w:tc><w:tcPr><w:tcW w:w="2400" w:type="dxa"/></w:tcPr>'
                f'{_paragraph_xml(text, style="Normal")}</w:tc>'
            )
        row_xml.append(f"<w:tr>{''.join(cells)}</w:tr>")
    return (
        '<w:tbl><w:tblPr><w:tblW w:w="0" w:type="auto"/>'
        '<w:tblBorders><w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/></w:tblBorders>'
        f"</w:tblPr><w:tblGrid>{grid}</w:tblGrid>{''.join(row_xml)}</w:tbl>"
    )


def _run_xml(text: str) -> str:
    preserve = ' xml:space="preserve"' if text.startswith(" ") or text.endswith(" ") else ""
    return f"<w:r><w:t{preserve}>{escape(text)}</w:t></w:r>"


def _content_types_xml() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
        '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
        "</Types>"
    )


def _root_relationships_xml() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>'
        "</Relationships>"
    )


def _document_relationships_xml() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
        "</Relationships>"
    )


def _core_properties_xml(title: str) -> str:
    created = datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace("+00:00", "Z")
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
        'xmlns:dc="http://purl.org/dc/elements/1.1/" '
        'xmlns:dcterms="http://purl.org/dc/terms/" '
        'xmlns:dcmitype="http://purl.org/dc/dcmitype/" '
        'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
        f"<dc:title>{escape(title)}</dc:title>"
        "<dc:creator>RegPilot</dc:creator>"
        f'<dcterms:created xsi:type="dcterms:W3CDTF">{created}</dcterms:created>'
        "</cp:coreProperties>"
    )


def _styles_xml() -> str:
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:rPr><w:rFonts w:ascii="Microsoft YaHei" w:eastAsia="Microsoft YaHei"/><w:sz w:val="22"/></w:rPr></w:style>'
        '<w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:basedOn w:val="Normal"/><w:pPr><w:spacing w:after="240"/></w:pPr><w:rPr><w:b/><w:rFonts w:ascii="Microsoft YaHei" w:eastAsia="Microsoft YaHei"/><w:sz w:val="36"/></w:rPr></w:style>'
        '<w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:pPr><w:spacing w:before="360" w:after="160"/></w:pPr><w:rPr><w:b/><w:sz w:val="30"/></w:rPr></w:style>'
        '<w:style w:type="paragraph" w:styleId="Heading2"><w:name w:val="heading 2"/><w:basedOn w:val="Normal"/><w:pPr><w:spacing w:before="280" w:after="120"/></w:pPr><w:rPr><w:b/><w:sz w:val="26"/></w:rPr></w:style>'
        '<w:style w:type="paragraph" w:styleId="Heading3"><w:name w:val="heading 3"/><w:basedOn w:val="Normal"/><w:pPr><w:spacing w:before="220" w:after="100"/></w:pPr><w:rPr><w:b/><w:sz w:val="24"/></w:rPr></w:style>'
        '<w:style w:type="paragraph" w:styleId="ListParagraph"><w:name w:val="List Paragraph"/><w:basedOn w:val="Normal"/><w:pPr><w:ind w:left="420" w:hanging="240"/></w:pPr></w:style>'
        "</w:styles>"
    )
